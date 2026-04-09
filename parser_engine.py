import os
import io
import re
import json
import asyncio
from groq import Groq
from pathlib import Path


# ── Text Extractors ────────────────────────────────────────────────────────────

def extract_from_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)


def extract_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


EXTRACTOR_MAP = {
    "txt": extract_from_txt,
    "pdf": extract_from_pdf,
    "docx": extract_from_docx,
    "doc": extract_from_docx,
}


# ── Groq-Powered Parser ────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an elite resume parsing engine. Extract ALL information from resumes with maximum accuracy.

Return ONLY a valid JSON object (no markdown, no explanations) with this exact schema:

{
  "personal_info": {
    "full_name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "github": "",
    "portfolio": "",
    "other_links": []
  },
  "professional_summary": "",
  "work_experience": [
    {
      "company": "",
      "title": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "responsibilities": [],
      "achievements": []
    }
  ],
  "education": [
    {
      "institution": "",
      "degree": "",
      "field_of_study": "",
      "start_date": "",
      "end_date": "",
      "gpa": "",
      "honors": "",
      "relevant_courses": []
    }
  ],
  "skills": {
    "technical": [],
    "soft": [],
    "languages": [],
    "tools": [],
    "frameworks": [],
    "databases": [],
    "cloud": [],
    "other": []
  },
  "certifications": [
    {
      "name": "",
      "issuer": "",
      "date": "",
      "expiry": "",
      "credential_id": ""
    }
  ],
  "projects": [
    {
      "name": "",
      "description": "",
      "technologies": [],
      "url": "",
      "start_date": "",
      "end_date": ""
    }
  ],
  "publications": [
    {
      "title": "",
      "publisher": "",
      "date": "",
      "url": ""
    }
  ],
  "awards": [
    {
      "title": "",
      "issuer": "",
      "date": "",
      "description": ""
    }
  ],
  "volunteer_experience": [
    {
      "organization": "",
      "role": "",
      "start_date": "",
      "end_date": "",
      "description": ""
    }
  ],
  "languages_spoken": [
    {
      "language": "",
      "proficiency": ""
    }
  ],
  "interests": [],
  "references": [],
  "metadata": {
    "total_experience_years": 0,
    "seniority_level": "",
    "primary_domain": "",
    "key_technologies": [],
    "resume_score": 0,
    "parsing_confidence": 0
  }
}

Rules:
1. Extract EVERY piece of information — leave nothing behind
2. If a field is not present, use null or empty string/array — never omit keys
3. Separate responsibilities (duties) from achievements (quantified results with numbers/%)
4. For skills, categorize precisely — a Python framework goes under "frameworks", not "technical"
5. Calculate total_experience_years from all work history (exclude overlaps)
6. Assign seniority_level: intern / junior / mid / senior / lead / principal / executive
7. Set resume_score 0-100 based on completeness and quality
8. Set parsing_confidence 0-100 based on how clearly the text was structured
9. NEVER hallucinate — only include data explicitly present in the resume
10. Extract dates in format: "Month YYYY" or "YYYY" if month unknown"""


class ResumeParser:
    def __init__(self):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY environment variable not set")
        self.client = Groq(api_key=api_key)
        self.model = "llama-3.3-70b-versatile"   # Best accuracy on Groq

    def _extract_text(self, file_bytes: bytes, ext: str) -> str:
        extractor = EXTRACTOR_MAP.get(ext)
        if not extractor:
            raise ValueError(f"No extractor for .{ext}")
        raw = extractor(file_bytes)
        # Normalize whitespace while preserving structure
        lines = [line.strip() for line in raw.splitlines()]
        cleaned = "\n".join(line for line in lines if line)
        return cleaned

    def _clean_json_response(self, text: str) -> str:
        """Strip any markdown fences or preamble the model might add."""
        # Remove ```json ... ``` fences
        text = re.sub(r"```(?:json)?", "", text).strip()
        # Find first { and last }
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1:
            raise ValueError("No JSON object found in model response")
        return text[start : end + 1]

    def _call_groq(self, resume_text: str) -> dict:
        # Trim to model context window (approx 10k tokens = ~40k chars is safe)
        if len(resume_text) > 40_000:
            resume_text = resume_text[:40_000]

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"Parse this resume:\n\n{resume_text}"},
            ],
            temperature=0.0,   # Deterministic for accuracy
            max_tokens=4096,
        )

        raw = response.choices[0].message.content
        cleaned = self._clean_json_response(raw)
        return json.loads(cleaned)

    async def parse(self, file_bytes: bytes, ext: str, filename: str) -> dict:
        loop = asyncio.get_event_loop()

        # Text extraction (sync, run in executor for non-blocking)
        resume_text = await loop.run_in_executor(
            None, self._extract_text, file_bytes, ext
        )

        if not resume_text.strip():
            raise ValueError("Could not extract any text from the file. The file may be scanned/image-based.")

        # Groq API call (sync SDK, run in executor)
        parsed = await loop.run_in_executor(None, self._call_groq, resume_text)

        # Attach raw text length as a debug helper
        parsed["metadata"]["char_count"] = len(resume_text)
        parsed["metadata"]["filename"] = filename

        return parsed
